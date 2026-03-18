"""
Export functionality — generates reports in Excel, PDF, and Word formats.
"""

import os
from datetime import datetime
from .config import EXPORTS_DIR
from .timer import get_sessions_in_range


def _ensure_exports_dir():
    os.makedirs(EXPORTS_DIR, exist_ok=True)


def _format_time(seconds: float, use_decimal: bool) -> str:
    """Format seconds into either HH:MM:SS or decimal hours based on user preference."""
    seconds = seconds or 0
    if use_decimal:
        return f"{seconds / 3600:.2f}h"
    
    seconds = max(0, int(seconds))
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    secs = seconds % 60
    
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"


def _get_week_description(date_str: str) -> str:
    """Returns a string like 'Week 1 of Jan'."""
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        month_name = dt.strftime("%b")
        week_num = (dt.day - 1) // 7 + 1
        return f"Week {week_num} of {month_name}"
    except:
        return "Unknown Week"


def _get_detailed_desc(s: dict) -> str:
    sub = s.get("subcategory_name") or ""
    notes = s.get("notes") or ""
    if sub and sub != "General" and notes:
        return f"{sub} — {notes}"
    elif notes:
        return notes
    else:
        return sub or "N/A"


def _build_report_data(start_date: str, end_date: str, project_id: str = "all", simple_format: bool = False, prepared_for: str = "") -> dict:
    """Gather all data needed for a report."""
    sessions = get_sessions_in_range(start_date, end_date)

    by_project = {}
    total_active = 0.0
    total_break = 0.0
    total_sessions_count = 0
    filtered_sessions = []

    for s in sessions:
        s_proj_id = s.get("project_id")
        if project_id != "all" and str(s_proj_id) != str(project_id):
            continue

        filtered_sessions.append(s)
        total_sessions_count += 1
        cat = s.get("category_name") or "Uncategorized"
        proj = s.get("project_name") or "Unassigned"

        if proj not in by_project:
            by_project[proj] = {
                "active_seconds": 0.0,
                "break_seconds": 0.0,
                "session_count": 0,
                "categories": {}
            }
        
        by_project[proj]["active_seconds"] += s["total_active_seconds"] or 0
        by_project[proj]["break_seconds"] += s["total_break_seconds"] or 0
        by_project[proj]["session_count"] += 1

        if cat not in by_project[proj]["categories"]:
            by_project[proj]["categories"][cat] = {
                "active_seconds": 0.0,
                "break_seconds": 0.0,
                "session_count": 0,
                "sessions": []
            }
            
        by_project[proj]["categories"][cat]["active_seconds"] += s["total_active_seconds"] or 0
        by_project[proj]["categories"][cat]["break_seconds"] += s["total_break_seconds"] or 0
        by_project[proj]["categories"][cat]["session_count"] += 1
        by_project[proj]["categories"][cat]["sessions"].append(s)

        total_active += s["total_active_seconds"] or 0
        total_break += s["total_break_seconds"] or 0

    # If simple format, group sessions by week
    simple_sessions = []
    if simple_format:
        # Sort by date
        filtered_sessions.sort(key=lambda x: x.get("date", ""))
        current_week = None
        for s in filtered_sessions:
            wd = _get_week_description(s.get("date", ""))
            if wd != current_week:
                current_week = wd
                simple_sessions.append({"is_header": True, "title": wd})
            simple_sessions.append(s)

    return {
        "by_project": by_project,
        "sessions_list": filtered_sessions,
        "simple_sessions": simple_sessions,
        "total_active": total_active,
        "total_break": total_break,
        "total_sessions": total_sessions_count,
        "start_date": start_date,
        "end_date": end_date,
        "prepared_for": prepared_for,
        "generated_at": datetime.now().isoformat(),
    }


# ─────────────────────────── EXCEL EXPORT ───────────────────────────


def export_excel(start_date: str, end_date: str, project_id: str = "all", decimal_format: bool = False, simple_format: bool = False, prepared_for: str = "") -> str:
    """Export report as a styled Excel workbook. Returns file path."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    _ensure_exports_dir()
    data = _build_report_data(start_date, end_date, project_id, simple_format)
    wb = Workbook()

    header_font = Font(name="Calibri", bold=True, size=12, color="FFFFFF")
    header_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
    title_font = Font(name="Calibri", bold=True, size=16, color="2F5496")
    subtitle_font = Font(name="Calibri", size=11, color="666666")
    data_font = Font(name="Calibri", size=11)
    border = Border(
        left=Side(style="thin", color="D9D9D9"),
        right=Side(style="thin", color="D9D9D9"),
        top=Side(style="thin", color="D9D9D9"),
        bottom=Side(style="thin", color="D9D9D9"),
    )
    
    ws = wb.active
    ws.title = "Report"
    
    if simple_format:
        # No totals, date first
        start_row = 1
        headers = ["Date", "Project", "Description", "Detailed Description", "Time(hours)"]
        for i, h in enumerate(headers, 1):
            cell = ws.cell(row=start_row, column=i, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = border
            
        row = start_row + 1
        for s in data["simple_sessions"]:
            if s.get("is_header"):
                ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=5)
                cell = ws.cell(row=row, column=1, value=s["title"])
                cell.font = Font(name="Calibri", bold=True, size=11)
                cell.fill = PatternFill(start_color="E9ECEF", end_color="E9ECEF", fill_type="solid")
                cell.alignment = Alignment(horizontal="left")
                for c in range(1, 6):
                    ws.cell(row=row, column=c).border = border
                row += 1
                continue

            ws.cell(row=row, column=1, value=s.get("date", "—")).font = data_font
            ws.cell(row=row, column=2, value=s.get("project_name") or "Unassigned").font = data_font
            ws.cell(row=row, column=3, value=s.get("category_name") or "Uncategorized").font = data_font
            ws.cell(row=row, column=4, value=_get_detailed_desc(s)).font = data_font
            
            active_h = (s.get("total_active_seconds") or 0) / 3600.0
            ws.cell(row=row, column=5, value=f"{active_h:.2f}h").font = data_font
            
            for c in range(1, 6):
                ws.cell(row=row, column=c).border = border
            row += 1

        ws.column_dimensions["A"].width = 15
        ws.column_dimensions["B"].width = 25
        ws.column_dimensions["C"].width = 25
        ws.column_dimensions["D"].width = 50
        ws.column_dimensions["E"].width = 15

    else:
        # Standard format with totals
        ws.merge_cells("A1:G1")
        ws["A1"] = "Productivity Report"
        ws["A1"].font = title_font

        ws.merge_cells("A2:G2")
        subtitle = f"{start_date} to {end_date} | Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        if prepared_for:
            subtitle = f"Prepared for: {prepared_for} | {subtitle}"
        ws["A2"] = subtitle
        ws["A2"].font = subtitle_font

        ws["A4"] = "Total Active Time"
        ws["B4"] = _format_time(data["total_active"], decimal_format)
        ws["A5"] = "Total Break Time"
        ws["B5"] = _format_time(data["total_break"], decimal_format)
        ws["A6"] = "Total Sessions"
        ws["B6"] = data["total_sessions"]
        for r in range(4, 7):
            ws.cell(row=r, column=1).font = Font(name="Calibri", bold=True, size=11)
            ws.cell(row=r, column=2).font = data_font

        start_row = 8
        headers = ["Project / Description / Date", "Start", "End", "Subcategory", "Active Time", "Break Time", "Detailed Description"]
        for i, h in enumerate(headers, 1):
            cell = ws.cell(row=start_row, column=i, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = border
            
        row = start_row + 1
        
        for proj_name, proj_data in sorted(data["by_project"].items(), key=lambda x: x[1]["active_seconds"], reverse=True):
            # Project Row
            ws.cell(row=row, column=1, value=proj_name).font = Font(name="Calibri", bold=True, size=12)
            ws.cell(row=row, column=5, value=_format_time(proj_data["active_seconds"], decimal_format)).font = Font(name="Calibri", bold=True, size=11)
            ws.cell(row=row, column=6, value=_format_time(proj_data["break_seconds"], decimal_format)).font = Font(name="Calibri", bold=True, size=11)
            
            for c in range(1, 8):
                ws.cell(row=row, column=c).fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
                ws.cell(row=row, column=c).border = border
            row += 1
            
            for cat_name, cat_data in sorted(proj_data["categories"].items(), key=lambda x: x[1]["active_seconds"], reverse=True):
                # Category Row
                ws.cell(row=row, column=1, value=f"  {cat_name}").font = Font(name="Calibri", bold=True, size=11)
                ws.cell(row=row, column=5, value=_format_time(cat_data["active_seconds"], decimal_format)).font = Font(name="Calibri", bold=True, size=11)
                ws.cell(row=row, column=6, value=_format_time(cat_data["break_seconds"], decimal_format)).font = Font(name="Calibri", bold=True, size=11)
                
                for c in range(1, 8):
                    ws.cell(row=row, column=c).fill = PatternFill(start_color="F2F7FB", end_color="F2F7FB", fill_type="solid")
                    ws.cell(row=row, column=c).border = border
                row += 1
                
                for s in cat_data["sessions"]:
                    start_time_str = s.get("start_time")
                    start_val = datetime.fromisoformat(start_time_str).strftime("%H:%M:%S") if start_time_str else "—"
                    end_time_str = s.get("end_time")
                    end_val = datetime.fromisoformat(end_time_str).strftime("%H:%M:%S") if end_time_str else "—"
                    
                    ws.cell(row=row, column=1, value=f"    {s.get('date', '—')}").font = data_font
                    ws.cell(row=row, column=2, value=start_val).font = data_font
                    ws.cell(row=row, column=3, value=end_val).font = data_font
                    ws.cell(row=row, column=4, value=s.get("subcategory_name") or "N/A").font = data_font
                    ws.cell(row=row, column=5, value=_format_time(s.get("total_active_seconds") or 0, decimal_format)).font = data_font
                    ws.cell(row=row, column=6, value=_format_time(s.get("total_break_seconds") or 0, decimal_format)).font = data_font
                    ws.cell(row=row, column=7, value=s.get("notes") or "").font = data_font
                    
                    for c in range(1, 8):
                        ws.cell(row=row, column=c).border = border
                    row += 1

        ws.column_dimensions["A"].width = 30
        ws.column_dimensions["B"].width = 12
        ws.column_dimensions["C"].width = 12
        ws.column_dimensions["D"].width = 20
        ws.column_dimensions["E"].width = 14
        ws.column_dimensions["F"].width = 14
        ws.column_dimensions["G"].width = 40

    filename = f"report_{start_date}_to_{end_date}.xlsx"
    filepath = os.path.join(EXPORTS_DIR, filename)
    wb.save(filepath)
    return filepath


# ─────────────────────────── PDF EXPORT ───────────────────────────


def export_pdf(start_date: str, end_date: str, project_id: str = "all", decimal_format: bool = False, simple_format: bool = False, prepared_for: str = "") -> str:
    """Export report as a professional PDF. Returns file path."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, mm
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, HRFlowable, KeepTogether
    )
    from reportlab.lib.enums import TA_CENTER

    _ensure_exports_dir()
    data = _build_report_data(start_date, end_date, project_id, simple_format)

    filename = f"report_{start_date}_to_{end_date}.pdf"
    filepath = os.path.join(EXPORTS_DIR, filename)

    doc = SimpleDocTemplate(
        filepath, pagesize=landscape(A4),
        rightMargin=15 * mm, leftMargin=15 * mm,
        topMargin=15 * mm, bottomMargin=15 * mm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportTitle", parent=styles["Title"],
        fontSize=22, textColor=colors.HexColor("#2F5496"),
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        "ReportSubtitle", parent=styles["Normal"],
        fontSize=11, textColor=colors.HexColor("#666666"),
        spaceAfter=20, alignment=TA_CENTER,
    )
    heading_style = ParagraphStyle(
        "SectionHeading", parent=styles["Heading2"],
        fontSize=14, textColor=colors.HexColor("#2F5496"),
        spaceBefore=16, spaceAfter=8,
    )
    normal_style = styles["Normal"]

    elements = []

    # Title & Totals only for standard format
    if not simple_format:
        elements.append(Paragraph("Productivity Report", title_style))
        subtitle_text = f"{start_date} to {end_date} &nbsp;|&nbsp; Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        if prepared_for:
            subtitle_text = f"Prepared for: {prepared_for} &nbsp;|&nbsp; {subtitle_text}"
        elements.append(Paragraph(subtitle_text, subtitle_style))
        elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#2F5496")))
        elements.append(Spacer(1, 12))

        # Overview stats
        overview_data = [
            ["Total Active Time:", _format_time(data["total_active"], decimal_format)],
            ["Total Break Time:", _format_time(data["total_break"], decimal_format)],
            ["Total Sessions:", str(data["total_sessions"])],
        ]
        overview_table = Table(overview_data, colWidths=[1.5 * inch, 2 * inch])
        overview_table.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 11),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#2F5496")),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
        ]))
        elements.append(overview_table)
        elements.append(Spacer(1, 16))

    # Main Table
    if simple_format:
        elements.append(Paragraph("Productivity Report (Simple)", title_style))
        elements.append(Paragraph(f"{start_date} to {end_date}", subtitle_style))
        
        headers = ["Date", "Project", "Description", "Detailed Description", "Time(hours)"]
        table_rows = [headers]
        col_widths = [1.0 * inch, 1.8 * inch, 1.8 * inch, 4.2 * inch, 1.0 * inch]
        
        row_styles = []
        row_idx = 1
        for s in data["simple_sessions"]:
            if s.get("is_header"):
                table_rows.append([s["title"], "", "", "", ""])
                row_styles.append(("BACKGROUND", (0, row_idx), (-1, row_idx), colors.HexColor("#E9ECEF")))
                row_styles.append(("FONTNAME", (0, row_idx), (-1, row_idx), "Helvetica-Bold"))
                row_styles.append(("SPAN", (0, row_idx), (-1, row_idx)))
                row_idx += 1
                continue

            notes_text = _get_detailed_desc(s).replace('\n', ' ')[:150]
            active_h = (s.get("total_active_seconds") or 0) / 3600.0
            table_rows.append([
                s.get("date", "—"),
                s.get("project_name") or "Unassigned",
                s.get("category_name") or "Uncategorized",
                Paragraph(notes_text, ParagraphStyle("Notes", parent=styles["Normal"], fontSize=9)),
                f"{active_h:.2f}h"
            ])
            if row_idx % 2 == 0:
                row_styles.append(("BACKGROUND", (0, row_idx), (-1, row_idx), colors.HexColor("#F2F7FB")))
            row_idx += 1

    else:
        elements.append(KeepTogether([
            Paragraph("Time by Project & Description", heading_style),
            Spacer(1, 4)
        ]))
        
        headers = ["Project / Description / Date", "Start", "End", "Subcat.", "Active", "Break", "Detailed Description"]
        table_rows = [headers]
        col_widths = [2.2 * inch, 0.7 * inch, 0.7 * inch, 1.2 * inch, 0.8 * inch, 0.8 * inch, 4.1 * inch]
        
        row_styles = []
        row_idx = 1
        for proj_name, proj_data in sorted(data["by_project"].items(), key=lambda x: x[1]["active_seconds"], reverse=True):
            table_rows.append([
                proj_name, "", "", "",
                _format_time(proj_data["active_seconds"], decimal_format),
                _format_time(proj_data["break_seconds"], decimal_format),
                ""
            ])
            row_styles.append(("BACKGROUND", (0, row_idx), (-1, row_idx), colors.HexColor("#D9E1F2")))
            row_styles.append(("FONTNAME", (0, row_idx), (-1, row_idx), "Helvetica-Bold"))
            row_idx += 1
            
            for cat_name, cat_data in sorted(proj_data["categories"].items(), key=lambda x: x[1]["active_seconds"], reverse=True):
                table_rows.append([
                    f"  {cat_name}", "", "", "",
                    _format_time(cat_data["active_seconds"], decimal_format),
                    _format_time(cat_data["break_seconds"], decimal_format),
                    ""
                ])
                row_styles.append(("BACKGROUND", (0, row_idx), (-1, row_idx), colors.HexColor("#F2F7FB")))
                row_styles.append(("FONTNAME", (0, row_idx), (-1, row_idx), "Helvetica-Bold"))
                row_idx += 1
                
                for s in cat_data["sessions"]:
                    start_time_str = s.get("start_time")
                    start_val = datetime.fromisoformat(start_time_str).strftime("%H:%M") if start_time_str else "—"
                    end_time_str = s.get("end_time")
                    end_val = datetime.fromisoformat(end_time_str).strftime("%H:%M") if end_time_str else "—"
                    notes_text = (s.get("notes") or "").replace('\n', ' ')[:100]
                    
                    table_rows.append([
                        f"    {s.get('date', '—')}",
                        start_val,
                        end_val,
                        s.get("subcategory_name") or "N/A",
                        _format_time(s.get("total_active_seconds") or 0, decimal_format),
                        _format_time(s.get("total_break_seconds") or 0, decimal_format),
                        Paragraph(notes_text, ParagraphStyle("Notes", parent=styles["Normal"], fontSize=9))
                    ])
                    row_idx += 1

    main_table = Table(table_rows, colWidths=col_widths, repeatRows=1)
    
    base_style = [
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F5496")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D9D9D9")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
    ]
    base_style.extend(row_styles)
    
    main_table.setStyle(TableStyle(base_style))
    elements.append(main_table)

    # Footer
    elements.append(Spacer(1, 20))
    elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#D9D9D9")))
    elements.append(Paragraph(
        f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        ParagraphStyle("Footer", parent=normal_style, fontSize=8, textColor=colors.HexColor("#999999"), alignment=TA_CENTER),
    ))

    doc.build(elements)
    return filepath


# ─────────────────────────── WORD EXPORT ───────────────────────────


def export_word(start_date: str, end_date: str, project_id: str = "all", decimal_format: bool = False, simple_format: bool = False, prepared_for: str = "") -> str:
    """Export report as a Word document. Returns file path."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT

    _ensure_exports_dir()
    data = _build_report_data(start_date, end_date, project_id, simple_format)

    doc = Document()

    # Switch to Landscape
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading("Productivity Report", level=0)
    title.runs[0].font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_text = f"{start_date} to {end_date}  |  Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    if prepared_for:
        subtitle_text = f"Prepared for: {prepared_for}  |  {subtitle_text}"
    run = subtitle.add_run(subtitle_text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # Totals only for standard
    if not simple_format:
        doc.add_heading("Overview", level=1)
        overview_table = doc.add_table(rows=3, cols=2)
        overview_data = [
            ("Total Active Time:", _format_time(data["total_active"], decimal_format)),
            ("Total Break Time:", _format_time(data["total_break"], decimal_format)),
            ("Total Sessions:", str(data["total_sessions"])),
        ]
        for i, (label, value) in enumerate(overview_data):
            overview_table.rows[i].cells[0].text = label
            overview_table.rows[i].cells[1].text = value
            overview_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()

    # Main Table
    if simple_format:
        doc.add_heading("Simple Session Log", level=1)
        headers = ["Date", "Project", "Description", "Detailed Description", "Time(hours)"]
        main_table = doc.add_table(rows=1 + len(data["simple_sessions"]), cols=len(headers))
        main_table.style = "Light Grid Accent 1"
        
        for i, h in enumerate(headers):
            main_table.cell(0, i).text = h
            main_table.cell(0, i).paragraphs[0].runs[0].font.bold = True
            
        row_idx = 1
        for s in data["simple_sessions"]:
            if s.get("is_header"):
                row_cell = main_table.cell(row_idx, 0)
                row_cell.text = s["title"]
                # Merge row
                for c in range(1, 5):
                    row_cell.merge(main_table.cell(row_idx, c))
                for p in row_cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                row_idx += 1
                continue

            main_table.cell(row_idx, 0).text = s.get("date", "—")
            main_table.cell(row_idx, 1).text = s.get("project_name") or "Unassigned"
            main_table.cell(row_idx, 2).text = s.get("category_name") or "Uncategorized"
            main_table.cell(row_idx, 3).text = _get_detailed_desc(s).replace('\n', ' ')
            
            active_h = (s.get("total_active_seconds") or 0) / 3600.0
            main_table.cell(row_idx, 4).text = f"{active_h:.2f}h"
            row_idx += 1

    else:
        doc.add_heading("Time by Project & Description", level=1)
        headers = ["Project / Description / Date", "Start", "End", "Subcat.", "Active", "Break", "Detailed Description"]
        
        num_rows = 1
        for proj_data in data["by_project"].values():
            num_rows += 1
            for cat_data in proj_data["categories"].values():
                num_rows += 1
                num_rows += len(cat_data["sessions"])
                
        main_table = doc.add_table(rows=num_rows, cols=len(headers))
        main_table.style = "Light Grid Accent 1"
        
        # Headers
        for i, h in enumerate(headers):
            main_table.cell(0, i).text = h
            main_table.cell(0, i).paragraphs[0].runs[0].font.bold = True
            
        row_idx = 1
        for proj_name, proj_data in sorted(data["by_project"].items(), key=lambda x: x[1]["active_seconds"], reverse=True):
            main_table.cell(row_idx, 0).text = proj_name
            main_table.cell(row_idx, 4).text = _format_time(proj_data["active_seconds"], decimal_format)
            main_table.cell(row_idx, 5).text = _format_time(proj_data["break_seconds"], decimal_format)
            
            for c in range(len(headers)):
                for p in main_table.cell(row_idx, c).paragraphs:
                    for r in p.runs:
                        r.font.bold = True
            row_idx += 1
            
            for cat_name, cat_data in sorted(proj_data["categories"].items(), key=lambda x: x[1]["active_seconds"], reverse=True):
                main_table.cell(row_idx, 0).text = f"  {cat_name}"
                main_table.cell(row_idx, 4).text = _format_time(cat_data["active_seconds"], decimal_format)
                main_table.cell(row_idx, 5).text = _format_time(cat_data["break_seconds"], decimal_format)
                
                for c in range(len(headers)):
                    for p in main_table.cell(row_idx, c).paragraphs:
                        for r in p.runs:
                            r.font.bold = True
                row_idx += 1
                
                for s in cat_data["sessions"]:
                    start_time_str = s.get("start_time")
                    start_val = datetime.fromisoformat(start_time_str).strftime("%H:%M") if start_time_str else "—"
                    end_time_str = s.get("end_time")
                    end_val = datetime.fromisoformat(end_time_str).strftime("%H:%M") if end_time_str else "—"
                    
                    main_table.cell(row_idx, 0).text = f"    {s.get('date', '—')}"
                    main_table.cell(row_idx, 1).text = start_val
                    main_table.cell(row_idx, 2).text = end_val
                    main_table.cell(row_idx, 3).text = s.get("subcategory_name") or "N/A"
                    main_table.cell(row_idx, 4).text = _format_time(s.get("total_active_seconds") or 0, decimal_format)
                    main_table.cell(row_idx, 5).text = _format_time(s.get("total_break_seconds") or 0, decimal_format)
                    main_table.cell(row_idx, 6).text = (s.get("notes") or "").replace('\n', ' ')
                    row_idx += 1

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    filename = f"report_{start_date}_to_{end_date}.docx"
    filepath = os.path.join(EXPORTS_DIR, filename)
    doc.save(filepath)
    return filepath


# ─────────────────────────── DISPATCHER ───────────────────────────


def export_report(fmt: str, start_date: str, end_date: str, project_id: str = "all", decimal_format: bool = False, simple_format: bool = False, prepared_for: str = "") -> str:
    """Export a report in the given format. Returns file path."""
    exporters = {
        "excel": export_excel,
        "pdf": export_pdf,
        "word": export_word,
    }
    if fmt not in exporters:
        raise ValueError(f"Unknown format '{fmt}'. Choose from: {', '.join(exporters.keys())}")
    return exporters[fmt](start_date, end_date, project_id, decimal_format, simple_format, prepared_for)
